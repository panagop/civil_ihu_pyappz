from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def set_a5_page(doc):
    section = doc.sections[0]
    section.page_width = Cm(14.8)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)


def add_cover(doc):
    # Based on Scan Page 7
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ΔΙΕΘΝΕΣ ΠΑΝΕΠΙΣΤΗΜΙΟ ΤΗΣ ΕΛΛΑΔΟΣ\n(Σέρρες)")
    run.bold = True
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\nΣΧΟΛΗ ΜΗΧΑΝΙΚΩΝ\n(ΠΡΩΗΝ ΣΧΟΛΗ ΤΕΧΝΟΛΟΓΙΚΩΝ ΕΦΑΡΜΟΓΩΝ)")
    run.bold = True
    run.font.size = Pt(11)

    # Spacer
    doc.add_paragraph("\n\n\n")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ΒΙΒΛΙΑΡΙΟ\nΠΡΑΚΤΙΚΗΣ ΑΣΚΗΣΗΣ")
    run.bold = True
    run.font.size = Pt(24)

    # Spacer
    doc.add_paragraph("\n\n\n")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "ΤΜΗΜΑ\nΜΗΧΑΝΙΚΩΝ ΠΛΗΡΟΦΟΡΙΚΗΣ,\nΥΠΟΛΟΓΙΣΤΩΝ\nΚΑΙ ΤΗΛΕΠΙΚΟΙΝΩΝΙΩΝ")
    run.bold = True
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("(πρώην Τμήμα Μηχανικών Πληροφορικής Τ.Ε.)")
    run.italic = True
    run.font.size = Pt(10)

    doc.add_page_break()


def add_inner_title(doc):
    # Based on Scan Page 1
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Βιβλιάριο Πρακτικής Άσκησης\nκαι Ημερολόγιο")
    run.bold = True
    run.font.size = Pt(16)

    doc.add_paragraph("\n")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ΣΧΟΛΗ ΜΗΧΑΝΙΚΩΝ\n(ΠΡΩΗΝ ΣΧΟΛΗ ΤΕΧΝΟΛΟΓΙΚΩΝ ΕΦΑΡΜΟΓΩΝ)")
    run.bold = True
    run.font.size = Pt(10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "ΤΜΗΜΑ\nΜΗΧΑΝΙΚΩΝ ΠΛΗΡΟΦΟΡΙΚΗΣ,\nΥΠΟΛΟΓΙΣΤΩΝ\nΚΑΙ ΤΗΛΕΠΙΚΟΙΝΩΝΙΩΝ\n(πρώην Τμήμα Μηχανικών Πληροφορικής Τ.Ε.)")
    run.bold = True
    run.font.size = Pt(11)

    # Footer number 1
    p = doc.add_paragraph("\n\n\n\n1")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()


def add_student_details(doc):
    # Based on Scan Page 2
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ΣΤΟΙΧΕΙΑ ΑΣΚΟΥΜΕΝΟΥ ΦΟΙΤΗΤΗ")
    run.bold = True
    run.font.size = Pt(14)
    run.underline = True

    fields = [
        "ΟΝΟΜΑ:",
        "ΕΠΩΝΥΜΟ:",
        "ΟΝΟΜΑ ΠΑΤΕΡΑ:",
        "ΕΤΟΣ ΓΕΝΝΗΣΗΣ:",
        "ΤΟΠΟΣ ΚΑΤΟΙΚΙΑΣ:",
        "Α.Μ. ΦΟΙΤΗΤΗ:",
        "ΦΟΡΕΑΣ ΑΠΑΣΧΟΛΗΣΗΣ:",
        "ΕΝΑΡΞΗ ΠΡΑΚΤΙΚΗΣ ΑΣΚΗΣΗΣ:",
        "ΛΗΞΗ ΠΡΑΚΤΙΚΗΣ ΑΣΚΗΣΗΣ:",
        "ΤΗΛΕΦΩΝΟ:"
    ]

    for field in fields:
        p = doc.add_paragraph()
        run = p.add_run(f"{field}")
        run.bold = True
        p.add_run(" " + "_" * 30)

    # Footer number 5
    p = doc.add_paragraph("\n5")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()


def add_form_1(doc):
    # Based on Scan Page 3
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("ΕΝΤΥΠΟ 1")
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ΕΒΔΟΜΑΔΙΑΙΟ ΦΥΛΛΟ ΠΡΑΚΤΙΚΗΣ ΑΣΚΗΣΗΣ")
    run.bold = True
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.add_run("ΕΒΔΟΜΑΔΑ ΠΡΑΚΤΙΚΗΣ ΑΣΚΗΣΗΣ Η .................")

    p = doc.add_paragraph()
    p.add_run("ΑΠΟ ........................... MEXPI ...........................")

    p = doc.add_paragraph()
    run = p.add_run("ΠΕΡΙΛΗΨΗ ΕΡΓΑΣΙΩΝ ΠΟΥ ΕΚΤΕΛΕΣΤΗΚΑΝ")
    run.bold = True

    # Create a table for the input box feel
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    cell.text = "\n\n\n\n\n"  # Space for writing

    p = doc.add_paragraph()
    run = p.add_run("\nΠΑΡΑΤΗΡΗΣΕΙΣ ΕΠΟΠΤΗ ΦΟΡΕΑ")
    run.bold = True

    table2 = doc.add_table(rows=1, cols=1)
    table2.style = 'Table Grid'
    cell2 = table2.cell(0, 0)
    cell2.text = "\n\n\n"

    p = doc.add_paragraph("\n")
    p.add_run("ΥΠΟΓΡΑΦΗ ΑΣΚΟΥΜΕΝΟΥ\t\tΥΠΟΓΡΑΦΗ ΣΦΡΑΓΙΔΑ\n")
    p.add_run("\t\t\t\t\tΕΠΟΠΤΗ ΦΟΡΕΑ\n\n")
    p.add_run(
        "..........................................\t\t..........................................")

    p = doc.add_paragraph("ΗΜΕΡΟΜΗΝΙΑ .........................")

    # Footer number 7
    p = doc.add_paragraph("7")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()


def add_form_2(doc):
    # Based on Scan Page 4
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("ΕΝΤΥΠΟ 2")
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "ΣΗΜΕΙΩΣΕΙΣ / ΠΑΡΑΤΗΡΗΣΕΙΣ\nΑΣΚΟΥΜΕΝΟΥ/ΑΣΚΟΥΜΕΝΗΣ ΦΟΙΤΗΤΗ/ΦΟΙΤΗΤΡΙΑΣ")
    run.bold = True
    run.font.size = Pt(12)

    # Lines for writing
    for _ in range(15):
        doc.add_paragraph("_" * 55)

    p = doc.add_paragraph("\n")
    p.add_run("ΥΠΟΓΡΑΦΗ: ..........................................")
    p = doc.add_paragraph()
    p.add_run("ΗΜΕΡΟΜΗΝΙΑ: ......................................")

    # Footer number 37
    p = doc.add_paragraph("37")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()


def add_form_3(doc):
    # Based on Scan Page 5
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("ΕΝΤΥΠΟ 3")
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "ΣΗΜΕΙΩΣΕΙΣ / ΠΑΡΑΤΗΡΗΣΕΙΣ ΕΠΟΠΤΗ ΦΟΡΕΑ\n- ΑΞΙΟΛΟΓΗΣΗ ΠΡΑΚΤΙΚΗΣ ΑΣΚΗΣΗΣ -")
    run.bold = True
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ΣΥΝΟΠΤΙΚΗ ΑΞΙΟΛΟΓΗΣΗ ΑΣΚΟΥΜΕΝΟΥ ΦΟΙΤΗΤΗ ΑΠΟ ΕΠΟΠΤΗ ΦΟΡΕΑ")
    run.italic = True
    run.font.size = Pt(10)

    # Table
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "ΙΚΑΝΟΤΗΤΕΣ ΑΣΚΟΥΜΕΝΟΥ"
    hdr_cells[1].text = "ΠΟΛΥ ΚΑΛΑ"
    hdr_cells[2].text = "ΚΑΛΑ"
    hdr_cells[3].text = "ΜΕΤΡΙΑ"

    criteria = [
        "Πρωτοβουλία - Υπευθυνότητα",
        "Ικανότητα Συνεργασίας",
        "Ποιοτική Απόδοση",
        "Ποσοτική Απόδοση",
        "Επιμέλεια-Ζήλος-Τήρηση Ωραρίου",
        "Συνολική Αξιολόγηση"
    ]

    for criterion in criteria:
        row_cells = table.add_row().cells
        row_cells[0].text = criterion

    p = doc.add_paragraph("\n")
    p.add_run("ΥΠΟΓΡΑΦΗ/ΣΦΡΑΓΙΔΑ ΕΠΟΠΤΗ ΦΟΡΕΑ\n\n........................................................................")

    p = doc.add_paragraph()
    p.add_run("ΗΜΕΡΟΜΗΝΙΑ: ......................................")

    # Footer number 38
    p = doc.add_paragraph("38")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()


def add_form_4(doc):
    # Based on Scan Page 6
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("ΕΝΤΥΠΟ 4")
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "ΣΗΜΕΙΩΣΕΙΣ / ΠΑΡΑΤΗΡΗΣΕΙΣ ΕΠΟΠΤΗ ΕΚΠΑΙΔΕΥΤΙΚΟΥ\n- ΑΞΙΟΛΟΓΗΣΗ ΠΡΑΚΤΙΚΗΣ ΑΣΚΗΣΗΣ -")
    run.bold = True
    run.font.size = Pt(11)

    # Lines for writing
    for _ in range(5):
        doc.add_paragraph("_" * 55)

    p = doc.add_paragraph("\n")
    p.add_run("ΟΝΟΜΑΤΕΠΩΝΥΜΟ ΑΚΑΔΗΜΑΪΚΟΥ ΕΠΟΠΤΗ:")
    p = doc.add_paragraph(
        "........................................................................................")

    p = doc.add_paragraph()
    p.add_run("ΗΜΕΡΟΜΗΝΙΑ: ......................  ΥΠΟΓΡΑΦΗ: ......................")

    doc.add_paragraph("_" * 30)  # Separator line

    p = doc.add_paragraph()
    p.add_run("ΟΝΟΜΑΤΕΠΩΝΥΜΟ ΠΡΟΕΔΡΟΥ ΤΜΗΜΑΤΟΣ:")
    p = doc.add_paragraph(
        "........................................................................................")

    p = doc.add_paragraph()
    p.add_run("ΗΜΕΡΟΜΗΝΙΑ: ......................  ΥΠΟΓΡΑΦΗ: ......................")

    p = doc.add_paragraph("\n")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ΕΓΚΡΙΝΕΤΑΙ  /  ΑΠΟΡΡΙΠΤΕΤΑΙ")
    run.bold = True
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("(Διαγράφετε ανάλογα)")
    p.italic = True

    # Footer number 39
    p = doc.add_paragraph("39")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


# Execute
doc = Document()
set_a5_page(doc)
add_cover(doc)
add_inner_title(doc)
add_student_details(doc)
add_form_1(doc)
add_form_2(doc)
add_form_3(doc)
add_form_4(doc)

file_path = "Internship_Booklet.docx"
doc.save(file_path)
