import pandas as pd
from pathlib import Path
import streamlit as st
from datetime import datetime, timedelta
from streamlit_calendar import calendar
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
from collections import defaultdict

st.set_page_config(
    layout="wide",
    page_title="Εβδομαδιαίο Πρόγραμμα Μαθημάτων",
    page_icon="📅",
)

# Χρώματα ανά εξάμηνο σπουδών
SEMESTER_COLORS = {
    1: '#E74C3C',  2: '#3498DB',  3: '#2ECC71',  4: '#F39C12',  5: '#9B59B6',
    6: '#1ABC9C',  7: '#E67E22',  8: '#34495E',  9: '#16A085',  10: '#D35400',
}

st.title("📅 Εβδομαδιαίο Πρόγραμμα Μαθημάτων")

# Επιλογή εξαμήνου
period_selection = st.radio(
    "Επιλέξτε εξάμηνο:",
    options=["Χειμερινό", "Εαρινό"],
    index=1,
    key="period_selection"
)

st.markdown(f"Έχετε επιλέξει: **{period_selection} Εξάμηνο**")

# Ρυθμίσεις αρχείου
INPUT_EXCEL = Path(__file__).parent.parent.parent / \
    "files" / "timetables" / "2025-2026.xlsm"


def load_data() -> pd.DataFrame:
    """Διαβάζει τα δεδομένα από το Excel."""

    # Έλεγχος ύπαρξης αρχείου
    if not INPUT_EXCEL.exists():
        st.error(f"❌ Το αρχείο {INPUT_EXCEL} δεν βρέθηκε!")
        st.info(f"Αναζητούμενη διαδρομή: {INPUT_EXCEL.absolute()}")
        st.stop()

    try:
        # Έλεγχος διαθέσιμων sheets
        excel_file = pd.ExcelFile(INPUT_EXCEL)
        available_sheets = excel_file.sheet_names

        # Προσδιορισμός sheet name με βάση το εξάμηνο
        sheet_name = 'timetable'  # Προσαρμόστε ανάλογα με τα πραγματικά ονόματα των sheets

        if sheet_name not in available_sheets:
            st.error(f"❌ Το sheet '{sheet_name}' δεν βρέθηκε στο αρχείο!")
            st.info(f"Διαθέσιμα sheets: {', '.join(available_sheets)}")
            st.stop()

        df = pd.read_excel(INPUT_EXCEL, sheet_name=sheet_name)

        # Βεβαιώσου ότι τα ονόματα στηλών ταιριάζουν με αυτά
        required_cols = [
            "course_id",
            "course_name",
            "class_name",
            "semester",
            "teaching_period",
            "instructors",
            "day",
            "start_time",
            "duration",
            "room",
            "notes",
        ]
        missing = [c for c in required_cols if c not in df.columns]
        if missing:
            raise ValueError(f"Λείπουν οι στήλες: {missing}")

        # Φιλτράρισμα με βάση το teaching_period
        df = df[df['teaching_period'] == period_selection]

        # room / course_id can mix numeric codes (e.g. 101) and strings (e.g. "ΔΟΜ704");
        # normalize to string so pyarrow doesn't infer int64 and fail on the strings.
        def _to_str(v):
            if pd.isna(v):
                return ""
            if isinstance(v, float) and v.is_integer():
                return str(int(v))
            return str(v)

        df["room"] = df["room"].apply(_to_str)
        df["course_id"] = df["course_id"].apply(_to_str)

        # Δημιουργία συνδυαστικής στήλης για καλύτερη αναγνώριση
        df['full_class_name'] = df.apply(
            lambda row: f"{row['course_name']} - {row['class_name']}"
            if pd.notna(row['class_name']) else str(row['course_name']),
            axis=1
        )

        # Μετατροπή start_time σε ώρα (integer) για υπολογισμούς
        df['start_hour'] = df['start_time'].apply(
            lambda x: x.hour if hasattr(x, 'hour') else int(x)
        )

        # Υπολογισμός ώρας λήξης και δημιουργία end_time
        df['end_hour'] = df['start_hour'] + df['duration']
        df['end_time'] = df.apply(
            lambda row: f"{int(row['end_hour'])}:00",
            axis=1
        )

    except Exception as e:
        st.error(f"❌ Σφάλμα κατά το άνοιγμα του αρχείου: {e}")
        st.stop()

    return df


def create_weekly_timetable_document(df: pd.DataFrame, period: str) -> bytes:
    """Δημιουργεί Word έγγραφο με εβδομαδιαίο πρόγραμμα μαθημάτων"""
    doc = Document()

    # Ρυθμίσεις σελίδας - landscape
    section = doc.sections[0]
    section.orientation = 1  # Landscape
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)

    # Τίτλος
    title = doc.add_heading(
        f'Εβδομαδιαίο Πρόγραμμα Μαθημάτων - {period} Εξάμηνο 2025-2026', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Ομαδοποίηση ανά εξάμηνο σπουδών
    semesters = sorted(df['semester'].unique())

    for sem_idx, semester in enumerate(semesters):
        df_sem = df[df['semester'] == semester]

        if df_sem.empty:
            continue

        # Επικεφαλίδα εξαμήνου
        sem_heading = doc.add_heading(f'Εξάμηνο {int(semester)}', level=1)

        # Όλες οι ώρες από 9:00 έως 21:00
        time_slots = list(range(9, 22))  # 9, 10, 11, ..., 21
        time_slots_str = [f"{h}:00" for h in time_slots]

        # Ημέρες της εβδομάδας
        day_names = ['Δευτέρα', 'Τρίτη', 'Τετάρτη', 'Πέμπτη', 'Παρασκευή']

        # Δημιουργία dictionary για συγκέντρωση μαθημάτων ανά κελί
        cell_classes = defaultdict(list)

        for day_idx, day_name in enumerate(day_names):
            df_day = df_sem[df_sem['day'] == day_name]

            for _, class_row in df_day.iterrows():
                start_time_str = str(class_row['start_time'])
                if ':' in start_time_str:
                    class_hour = int(start_time_str.split(':')[0])
                else:
                    try:
                        class_hour = int(float(start_time_str))
                    except:
                        continue

                # Βρες το σωστό time slot index
                try:
                    time_idx = time_slots.index(class_hour)
                except ValueError:
                    continue

                # Προσθήκη μαθήματος στο dictionary με key (time_idx, day_idx)
                cell_key = (time_idx, day_idx)
                cell_classes[cell_key].append({
                    'course': str(class_row['full_class_name']) if pd.notna(class_row['full_class_name']) else '',
                    'instructor': str(class_row['instructors']) if pd.notna(class_row['instructors']) else '',
                    'room': str(class_row['room']) if pd.notna(class_row['room']) else '',
                    'duration': int(class_row['duration']) if pd.notna(class_row['duration']) else 1
                })

        # Υπολογισμός μέγιστου αριθμού ταυτόχρονων μαθημάτων
        max_simultaneous = 1
        for classes in cell_classes.values():
            max_simultaneous = max(max_simultaneous, len(classes))

        # Δημιουργία πίνακα με επιπλέον στήλες για ταυτόχρονα μαθήματα
        total_cols = 1 + (len(day_names) * max_simultaneous)
        table = doc.add_table(rows=len(time_slots) + 1, cols=total_cols)
        table.style = 'Light Grid Accent 1'

        # Επικεφαλίδα - στήλη ώρας
        table.rows[0].cells[0].text = 'Ώρα'

        # Επικεφαλίδες ημερών - merge cells για κάθε ημέρα
        for day_idx, day_name in enumerate(day_names):
            start_col = 1 + (day_idx * max_simultaneous)
            cell = table.rows[0].cells[start_col]

            # Merge across all sub-columns for this day
            if max_simultaneous > 1:
                for sub_col in range(1, max_simultaneous):
                    cell.merge(table.rows[0].cells[start_col + sub_col])

            cell.text = day_name
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = 'Calibri'
                    run.font.bold = True
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(255, 255, 255)

        # Συμπλήρωση σειρών ωρών και μαθημάτων
        # Track which (time_idx, day_idx, cls_idx) have been processed
        processed_cells = {}

        for time_idx, time_slot in enumerate(time_slots_str):
            row_idx = time_idx + 1

            # Στήλη ώρας
            cell = table.rows[row_idx].cells[0]
            cell.text = time_slot
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = 'Calibri'
                    run.font.bold = True
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(255, 255, 255)

            # Κελιά μαθημάτων για κάθε ημέρα
            for day_idx in range(len(day_names)):
                cell_key = (time_idx, day_idx)
                classes = cell_classes.get(cell_key, [])

                start_col = 1 + (day_idx * max_simultaneous)

                if len(classes) <= 1:
                    # Check if already processed as part of a duration span
                    if (time_idx, day_idx, 0) in processed_cells:
                        continue

                    # 0 ή 1 μάθημα - merge across all sub-columns
                    if len(classes) == 1:
                        cls = classes[0]
                        duration = cls['duration']

                        # Mark this cell and future rows as processed
                        for dur in range(duration):
                            for sub_col_idx in range(max_simultaneous):
                                processed_cells[(
                                    time_idx + dur, day_idx, sub_col_idx)] = True

                        # Collect all cells to merge into one rectangular region
                        cells_to_merge = []
                        for dur_offset in range(duration):
                            if time_idx + dur_offset < len(time_slots):
                                for sub_col_offset in range(max_simultaneous):
                                    cells_to_merge.append(
                                        table.rows[row_idx + dur_offset].cells[start_col + sub_col_offset])

                        # Start with first cell
                        cell = cells_to_merge[0]

                        # Merge all cells in the rectangular region
                        for merge_cell in cells_to_merge[1:]:
                            if merge_cell != cell:  # Don't merge with itself
                                try:
                                    cell.merge(merge_cell)
                                except Exception:
                                    pass  # Skip if already merged

                        class_text = f"{cls['course']}\n{cls['instructor']}"
                        if cls['room']:
                            class_text += f"\n{cls['room']}"
                        cell.text = class_text

                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for run in paragraph.runs:
                                run.font.name = 'Calibri'
                                run.font.size = Pt(8)
                    else:
                        # Κενό κελί - merge horizontally across sub-columns
                        # Mark all sub-columns as processed to avoid double-processing
                        for sub_col_idx in range(max_simultaneous):
                            processed_cells[(
                                time_idx, day_idx, sub_col_idx)] = True

                        # Merge across all sub-columns for this row only
                        if max_simultaneous > 1:
                            cell = table.rows[row_idx].cells[start_col]
                            for sub_col_offset in range(1, max_simultaneous):
                                next_col_cell = table.rows[row_idx].cells[start_col + sub_col_offset]
                                try:
                                    cell.merge(next_col_cell)
                                except Exception:
                                    pass  # Skip if already merged
                else:
                    # Πολλαπλά μαθήματα - ξεχωριστό κελί για καθένα
                    for cls_idx, cls in enumerate(classes[:max_simultaneous]):
                        # Check if already processed
                        if (time_idx, day_idx, cls_idx) in processed_cells:
                            continue

                        col = start_col + cls_idx
                        duration = cls['duration']

                        # Mark as processed
                        for dur in range(duration):
                            processed_cells[(
                                time_idx + dur, day_idx, cls_idx)] = True

                        cell = table.rows[row_idx].cells[col]

                        # Merge vertically for duration
                        if duration > 1:
                            for dur_offset in range(1, duration):
                                if time_idx + dur_offset < len(time_slots):
                                    next_cell = table.rows[row_idx +
                                                           dur_offset].cells[col]
                                    try:
                                        cell.merge(next_cell)
                                    except Exception:
                                        pass  # Skip if already merged

                        class_text = f"{cls['course']}\n{cls['instructor']}"
                        if cls['room']:
                            class_text += f"\n{cls['room']}"
                        cell.text = class_text

                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for run in paragraph.runs:
                                run.font.name = 'Calibri'
                                run.font.size = Pt(8)

        # Εφαρμογή χρωμάτων
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        # Χρώμα επικεφαλίδας
        for col_idx in range(0, total_cols):
            cell = table.rows[0].cells[col_idx]
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), '4472C4')
            cell._element.get_or_add_tcPr().append(shading_elm)

        # Χρώματα κελιών - λευκό για κενά, E7EFF7 για μαθήματα
        for row_idx in range(1, len(time_slots) + 1):
            for col_idx in range(0, total_cols):
                cell = table.rows[row_idx].cells[col_idx]
                shading_elm = OxmlElement('w:shd')

                if col_idx == 0:
                    # Στήλη ωρών - σκούρο μπλε
                    shading_elm.set(qn('w:fill'), '4472C4')
                else:
                    # Έλεγχος αν το κελί έχει περιεχόμενο (μάθημα)
                    cell_text = cell.text.strip()
                    if cell_text:
                        # Κελί με μάθημα - ανοιχτό μπλε
                        shading_elm.set(qn('w:fill'), 'E7EFF7')
                    else:
                        # Κενό κελί - λευκό
                        shading_elm.set(qn('w:fill'), 'FFFFFF')

                cell._element.get_or_add_tcPr().append(shading_elm)

        # Προσαρμογή πλάτους στηλών και ύψους σειρών
        for row in table.rows:
            row.height = Inches(0.4)  # Fixed height for all rows
            row.cells[0].width = Inches(0.7)
            for i in range(1, len(day_names) * max_simultaneous + 1):
                row.cells[i].width = Inches(1.5)

        # Page break μεταξύ εξαμήνων (εκτός από το τελευταίο)
        if sem_idx < len(semesters) - 1:
            doc.add_page_break()

    # Αποθήκευση σε buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# Φόρτωση δεδομένων
try:
    df = load_data()

    # st.subheader(f"Πρόγραμμα {semester_selection} Εξαμήνου 2025-2026")

    # Tabs
    tab_calendar, tab_table, tab_rooms, tab_instructors, tab_export = st.tabs(
        ["Εβδομαδιαία Προβολή", "Πίνακας", "Αιθουσιολόγιο", "Ανά Καθηγητή", "Εξαγωγή Word"])

    with tab_table:
        # Εμφάνιση δεδομένων (με end_time)
        display_cols = ['course_id', 'course_name', 'class_name', 'full_class_name', 'semester',
                        'teaching_period', 'instructors', 'day', 'start_time', 'end_time',
                        'duration', 'room', 'notes']
        available_cols = [col for col in display_cols if col in df.columns]
        st.dataframe(df[available_cols], width='stretch')

    with tab_calendar:
        st.markdown("### Εβδομαδιαίο Πρόγραμμα")

        # Φίλτρο εξαμήνων σπουδών
        semesters_all = sorted(df["semester"].unique().tolist())
        semester_options = [f"Εξάμηνο {int(s)}" for s in semesters_all]

        selected_semester = st.selectbox(
            "Επιλέξτε εξάμηνο σπουδών:",
            options=semester_options,
            index=0,
            key="semester_filter"
        )

        # Φιλτράρισμα δεδομένων
        semester_num = int(selected_semester.split()[-1])
        df_filtered = df[df["semester"] == semester_num]

        # Safely handle potential None values, convert to string, and remove problematic characters
        def clean_text(value):
            if pd.notna(value):
                # Convert to string and remove newlines, quotes, backslashes
                text = str(value).replace('\n', ' ').replace('\r', ' ')
                text = text.replace('"', '').replace("'", '').replace('\\', '')
                return text.strip()
            return ""

        # Convert to calendar events
        calendar_events = []

        # Map Greek days to weekday numbers (0=Monday)
        day_map = {
            'Δευτέρα': 0, 'Τρίτη': 1, 'Τετάρτη': 2, 'Πέμπτη': 3, 'Παρασκευή': 4,
            'Σάββατο': 5, 'Κυριακή': 6
        }

        # Use a reference week (e.g., a week in January 2025)
        reference_date = datetime(2025, 1, 6)  # Monday, January 6, 2025

        for _, row in df_filtered.iterrows():
            try:
                if pd.notna(row['day']) and pd.notna(row['start_time']):
                    # Get the day of week
                    day_name = str(row['day']).strip()
                    weekday = day_map.get(day_name, 0)

                    # Calculate the date for this event
                    event_date = reference_date + timedelta(days=weekday)

                    # Extract hour from start_time
                    start_time_str = str(row['start_time'])
                    if ':' in start_time_str:
                        start_hour = int(start_time_str.split(':')[0])
                    else:
                        start_hour = int(float(start_time_str))

                    # Create start datetime
                    start_dt = event_date.replace(
                        hour=start_hour, minute=0, second=0)
                    start_str = start_dt.strftime("%Y-%m-%dT%H:%M:%S")

                    # Calculate end time based on duration
                    duration = int(row['duration']) if pd.notna(
                        row['duration']) else 1
                    end_dt = start_dt + timedelta(hours=duration)
                    end_str = end_dt.strftime("%Y-%m-%dT%H:%M:%S")

                    # Clean text
                    full_class_name = clean_text(row['full_class_name'])
                    instructors = clean_text(row['instructors'])
                    room = clean_text(row['room'])
                    semester = int(row['semester']) if pd.notna(
                        row['semester']) else 1

                    # Get color
                    color = SEMESTER_COLORS.get(semester, '#95A5A6')

                    # Create concise title
                    title_parts = [full_class_name]
                    if instructors:
                        title_parts.append(instructors)
                    if room:
                        title_parts.append(f'({room})')

                    event = {
                        "title": ' - '.join(title_parts),
                        "start": start_str,
                        "end": end_str,
                        "color": color
                    }
                    calendar_events.append(event)
            except Exception as e:
                # Skip rows with errors
                continue

        st.write(f"📚 Σύνολο μαθημάτων: {len(calendar_events)}")

        # CSS to hide dates and make it look generic
        st.markdown("""
        <style>
        /* Hide the date numbers in column headers */
        .fc-col-header-cell-cushion {
            font-size: 14px !important;
        }
        .fc-daygrid-day-number {
            display: none !important;
        }
        /* Hide the full date range in title */
        .fc-toolbar-title {
            display: none !important;
        }
        /* Style for cleaner look */
        .fc-toolbar-chunk:first-child {
            display: none !important;
        }
        /* Preserve whitespace and line breaks in event titles */
        .fc-event-title, .fc-event-title-container, .fc-timegrid-event-harness, .fc-event-main {
            white-space: pre-line !important;
        }
        .fc-timegrid-event {
            white-space: pre-line !important;
        }
        </style>
        """, unsafe_allow_html=True)

        # Calendar options
        calendar_options = {
            "initialView": "timeGridWeek",
            "initialDate": "2025-01-06",  # Start on Monday
            "headerToolbar": {
                "left": "",
                "center": "",
                "right": ""
            },
            "slotMinTime": "08:00:00",
            "slotMaxTime": "21:00:00",
            "allDaySlot": False,
            "height": 850,
            "locale": "el",
            "firstDay": 1,  # Monday
            "weekends": False,  # Hide weekends
            "navLinks": False,
            "editable": False,
            "selectable": False,
            "dayHeaderFormat": {"weekday": "long"},  # Show only day names
            "displayEventTime": False,  # Hide time in event boxes
        }

        # Create a unique key based on selected semester
        semester_num = int(selected_semester.split()[-1])
        calendar_key = f"timetable_sem_{semester_num}"

        # Initialize session state for calendar display
        if 'show_timetable_calendar' not in st.session_state:
            st.session_state.show_timetable_calendar = False

        # Button to show calendar
        if not st.session_state.show_timetable_calendar:
            if st.button("📅 Εμφάνιση Ημερολογίου", key="show_timetable_cal_btn"):
                st.session_state.show_timetable_calendar = True
                st.rerun()

        # Render calendar if button was clicked
        if st.session_state.show_timetable_calendar:
            if calendar_events:
                calendar_data = calendar(
                    events=calendar_events,
                    options=calendar_options,
                    key=calendar_key
                )
            else:
                st.info(
                    "Δεν υπάρχουν μαθήματα για εμφάνιση με τα επιλεγμένα φίλτρα.")

    with tab_rooms:
        st.markdown("### Αιθουσιολόγιο - Πρόγραμμα Αιθουσών")

        # Φίλτρο αιθουσών
        rooms_all = sorted(
            [str(r) for r in df["room"].unique() if pd.notna(r) and str(r).strip()])

        if not rooms_all:
            st.warning("⚠️ Δεν βρέθηκαν αίθουσες στα δεδομένα.")
        else:
            selected_room = st.selectbox(
                "Επιλέξτε αίθουσα:",
                options=rooms_all,
                index=0,
                key="room_filter"
            )

            # Φιλτράρισμα δεδομένων ανά αίθουσα (convert to string for comparison)
            df_filtered_room = df[df["room"].astype(str) == selected_room]

            # Safely handle potential None values, convert to string, and remove problematic characters
            def clean_text_room(value):
                if pd.notna(value):
                    # Convert to string and remove newlines, quotes, backslashes
                    text = str(value).replace('\n', ' ').replace('\r', ' ')
                    text = text.replace('"', '').replace(
                        "'", '').replace('\\', '')
                    return text.strip()
                return ""

            # Convert to calendar events
            calendar_events_room = []

            # Map Greek days to weekday numbers (0=Monday)
            day_map_room = {
                'Δευτέρα': 0, 'Τρίτη': 1, 'Τετάρτη': 2, 'Πέμπτη': 3, 'Παρασκευή': 4,
                'Σάββατο': 5, 'Κυριακή': 6
            }

            # Use a reference week (e.g., a week in January 2025)
            reference_date_room = datetime(
                2025, 1, 6)  # Monday, January 6, 2025

            for _, row in df_filtered_room.iterrows():
                try:
                    if pd.notna(row['day']) and pd.notna(row['start_time']):
                        # Get the day of week
                        day_name = str(row['day']).strip()
                        weekday = day_map_room.get(day_name, 0)

                        # Calculate the date for this event
                        event_date = reference_date_room + \
                            timedelta(days=weekday)

                        # Extract hour from start_time
                        start_time_str = str(row['start_time'])
                        if ':' in start_time_str:
                            start_hour = int(start_time_str.split(':')[0])
                        else:
                            start_hour = int(float(start_time_str))

                        # Create start datetime
                        start_dt = event_date.replace(
                            hour=start_hour, minute=0, second=0)
                        start_str = start_dt.strftime("%Y-%m-%dT%H:%M:%S")

                        # Calculate end time based on duration
                        duration = int(row['duration']) if pd.notna(
                            row['duration']) else 1
                        end_dt = start_dt + timedelta(hours=duration)
                        end_str = end_dt.strftime("%Y-%m-%dT%H:%M:%S")

                        # Clean text
                        full_class_name = clean_text_room(
                            row['full_class_name'])
                        instructors = clean_text_room(row['instructors'])
                        semester = int(row['semester']) if pd.notna(
                            row['semester']) else 1

                        # Get color based on semester
                        color = SEMESTER_COLORS.get(semester, '#95A5A6')

                        # Create concise title (include semester info)
                        title_parts = [full_class_name]
                        if instructors:
                            title_parts.append(instructors)
                        title_parts.append(f'Εξ.{semester}')

                        event = {
                            "title": ' - '.join(title_parts),
                            "start": start_str,
                            "end": end_str,
                            "color": color
                        }
                        calendar_events_room.append(event)
                except Exception as e:
                    # Skip rows with errors
                    continue

            st.write(
                f"🏫 Σύνολο μαθημάτων στην αίθουσα {selected_room}: {len(calendar_events_room)}")

            # CSS to hide dates and make it look generic
            st.markdown("""
            <style>
            /* Hide the date numbers in column headers */
            .fc-col-header-cell-cushion {
                font-size: 14px !important;
            }
            .fc-daygrid-day-number {
                display: none !important;
            }
            /* Hide the full date range in title */
            .fc-toolbar-title {
                display: none !important;
            }
            /* Style for cleaner look */
            .fc-toolbar-chunk:first-child {
                display: none !important;
            }
            /* Preserve whitespace and line breaks in event titles */
            .fc-event-title, .fc-event-title-container, .fc-timegrid-event-harness, .fc-event-main {
                white-space: pre-line !important;
            }
            .fc-timegrid-event {
                white-space: pre-line !important;
            }
            </style>
            """, unsafe_allow_html=True)

            # Calendar options
            calendar_options_room = {
                "initialView": "timeGridWeek",
                "initialDate": "2025-01-06",  # Start on Monday
                "headerToolbar": {
                    "left": "",
                    "center": "",
                    "right": ""
                },
                "slotMinTime": "08:00:00",
                "slotMaxTime": "21:00:00",
                "allDaySlot": False,
                "height": 850,
                "locale": "el",
                "firstDay": 1,  # Monday
                "weekends": False,  # Hide weekends
                "navLinks": False,
                "editable": False,
                "selectable": False,
                "dayHeaderFormat": {"weekday": "long"},  # Show only day names
                "displayEventTime": False,  # Hide time in event boxes
            }

            # Create a unique key based on selected room
            calendar_key_room = f"room_timetable_{selected_room}"

            # Initialize session state for room calendar display
            if 'show_room_calendar' not in st.session_state:
                st.session_state.show_room_calendar = False

            # Button to show calendar
            if not st.session_state.show_room_calendar:
                if st.button("🏫 Εμφάνιση Αιθουσιολογίου", key="show_room_cal_btn"):
                    st.session_state.show_room_calendar = True
                    st.rerun()

            # Render calendar if button was clicked
            if st.session_state.show_room_calendar:
                if calendar_events_room:
                    calendar_data_room = calendar(
                        events=calendar_events_room,
                        options=calendar_options_room,
                        key=calendar_key_room
                    )
                else:
                    st.info("Δεν υπάρχουν μαθήματα στην επιλεγμένη αίθουσα.")

    with tab_instructors:
        st.markdown("### Μαθήματα Ανά Καθηγητή")

        # Δημιουργία λίστας με όλα τα μαθήματα ανά καθηγητή
        # Χειρισμός πολλαπλών καθηγητών στο ίδιο μάθημα
        instructor_classes = []

        for _, row in df.iterrows():
            if pd.notna(row['instructors']) and str(row['instructors']).strip():
                # Διαχωρισμός καθηγητών (με κόμμα, ερωτηματικό, ή άλλους διαχωριστές)
                instructors_list = str(row['instructors']).replace(
                    ';', ',').split(',')

                for instructor in instructors_list:
                    instructor = instructor.strip()
                    if instructor:  # Αγνόηση κενών εγγραφών
                        instructor_classes.append({
                            'Καθηγητής': instructor,
                            'Κωδικός': row['course_id'],
                            'Μάθημα': row['course_name'],
                            'Τμήμα': row['class_name'] if pd.notna(row['class_name']) else '',
                            'Εξάμηνο': int(row['semester']) if pd.notna(row['semester']) else '',
                            'Ημέρα': row['day'],
                            'Ώρα': row['start_time'],
                            'Διάρκεια': f"{int(row['duration'])}h" if pd.notna(row['duration']) else '',
                            'Αίθουσα': row['room'] if pd.notna(row['room']) else '',
                            'Παρατηρήσεις': row['notes'] if pd.notna(row['notes']) else ''
                        })

        if not instructor_classes:
            st.warning("⚠️ Δεν βρέθηκαν καθηγητές στα δεδομένα.")
        else:
            # Δημιουργία DataFrame
            df_instructors = pd.DataFrame(instructor_classes)

            # Ταξινόμηση κατά καθηγητή, εξάμηνο, ημέρα
            df_instructors = df_instructors.sort_values(
                by=['Καθηγητής', 'Εξάμηνο', 'Ημέρα', 'Ώρα'])

            # Φίλτρο καθηγητών
            all_instructors = sorted(
                df_instructors['Καθηγητής'].unique().tolist())

            col1, col2 = st.columns([2, 1])

            with col1:
                selected_instructor = st.selectbox(
                    "Επιλέξτε καθηγητή:",
                    options=['Όλοι'] + all_instructors,
                    index=0,
                    key="instructor_filter"
                )

            with col2:
                # Στατιστικά
                if selected_instructor == 'Όλοι':
                    st.metric("Σύνολο Καθηγητών", len(all_instructors))
                    st.metric("Σύνολο Μαθημάτων", len(df_instructors))
                else:
                    df_selected = df_instructors[df_instructors['Καθηγητής']
                                                 == selected_instructor]
                    st.metric("Μαθήματα", len(df_selected))
                    unique_courses = df_selected['Κωδικός'].nunique()
                    st.metric("Μοναδικά Μαθήματα", unique_courses)

            # Εμφάνιση πίνακα
            st.markdown("---")

            if selected_instructor == 'Όλοι':
                # Εμφάνιση ομαδοποιημένα ανά καθηγητή
                for instructor in all_instructors:
                    df_instr = df_instructors[df_instructors['Καθηγητής']
                                              == instructor]

                    # Υπολογισμός συνολικών ωρών (αφαίρεση 'h' από τη στήλη Διάρκεια)
                    total_hours = df_instr['Διάρκεια'].apply(lambda x: int(
                        str(x).replace('h', '')) if pd.notna(x) and str(x).strip() else 0).sum()

                    with st.expander(f"📚 {instructor} ({total_hours} ώρες)", expanded=False):
                        st.dataframe(
                            df_instr.drop(columns=['Καθηγητής']),
                            width='stretch',
                            hide_index=True
                        )
            else:
                # Εμφάνιση για επιλεγμένο καθηγητή
                df_selected = df_instructors[df_instructors['Καθηγητής']
                                             == selected_instructor]

                st.subheader(f"Μαθήματα: {selected_instructor}")
                st.dataframe(
                    df_selected.drop(columns=['Καθηγητής']),
                    width='stretch',
                    hide_index=True
                )

                # Ανάλυση ανά εξάμηνο
                st.markdown("#### Κατανομή ανά Εξάμηνο Σπουδών")
                semester_counts = df_selected.groupby(
                    'Εξάμηνο').size().reset_index(name='Πλήθος')

                col1, col2 = st.columns([1, 2])
                with col1:
                    st.dataframe(semester_counts, hide_index=True,
                                 width='stretch')
                with col2:
                    st.bar_chart(semester_counts.set_index('Εξάμηνο'))

    with tab_export:
        st.subheader("Εξαγωγή Εβδομαδιαίου Προγράμματος")
        st.markdown(
            "Δημιουργήστε αρχείο Word με το εβδομαδιαίο πρόγραμμα μαθημάτων για όλα τα εξάμηνα.")

        # Φίλτρα για εξαγωγή
        st.markdown("### Επιλογές Φιλτραρίσματος")

        col1, col2 = st.columns(2)

        with col1:
            semesters_export = sorted(df["semester"].unique().tolist())
            semester_options_export = [
                f"Εξάμηνο {int(s)}" for s in semesters_export]

            selected_export_semesters = st.multiselect(
                "Επιλέξτε εξάμηνα:",
                options=semester_options_export,
                default=semester_options_export,
                key="export_semester_filter"
            )

        # Φιλτράρισμα δεδομένων
        df_export = df.copy()

        if selected_export_semesters and len(selected_export_semesters) < len(semester_options_export):
            semester_nums = [int(s.split()[-1])
                             for s in selected_export_semesters]
            df_export = df_export[df_export["semester"].isin(semester_nums)]

        # Προεπισκόπηση
        st.markdown("### Προεπισκόπηση Δεδομένων")
        st.write(f"Σύνολο μαθημάτων προς εξαγωγή: {len(df_export)}")

        if not df_export.empty:
            st.dataframe(
                df_export[['day', 'start_time', 'semester', 'full_class_name',
                          'instructors', 'room', 'duration']].sort_values(by=['semester', 'day', 'start_time']),
                height=400
            )

            # Κουμπί λήψης
            st.markdown("### Λήψη Αρχείου")

            try:
                word_file = create_weekly_timetable_document(
                    df_export, period_selection)

                # Όνομα αρχείου
                filename = f"Προγραμμα_Μαθηματων_{period_selection}_2025-2026.docx"

                st.download_button(
                    label="📥 Λήψη Word Αρχείου",
                    data=word_file,
                    file_name=filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    help="Κατεβάστε το εβδομαδιαίο πρόγραμμα μαθημάτων σε μορφή Word"
                )

                st.success("✅ Το αρχείο είναι έτοιμο για λήψη!")

            except Exception as e:
                st.error(f"Σφάλμα κατά τη δημιουργία του αρχείου: {e}")
        else:
            st.warning("⚠️ Δεν υπάρχουν δεδομένα με τα επιλεγμένα φίλτρα.")

except Exception as e:
    st.error(f"Σφάλμα: {e}")
    st.info("Παρακαλώ ελέγξτε τη δομή του αρχείου Excel και τα ονόματα των sheets.")
