import pandas as pd
from pathlib import Path
from datetime import datetime, timedelta
import streamlit as st
from streamlit_calendar import calendar
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

exam_period = "Χειμερινό Εξάμηνο 2025-2026"

st.set_page_config(
    layout="wide",
    page_title=f"Πρόγραμμα Εξετάσεων - {exam_period}",
    page_icon="🗓️",

)

st.title(f"🗓️ Πρόγραμμα Εξετάσεων - {exam_period}")

program_selection: str = st.radio(
    "Επιλογή προγράμματος σπουδών:",
    options=["ΔΙΠΑΕ", "ΤΕΙ"],
    index=0,
    key="program_selection"
)

st.markdown(f"Έχετε επιλέξει το πρόγραμμα σπουδών: **{program_selection}**")

# ------------ ΡΥΘΜΙΣΕΙΣ ΧΡΗΣΤΗ ------------
# INPUT_EXCEL = "lessons-calendars.xlsm"   # το αρχείο όπου έχεις το sheet Data
INPUT_SHEET = program_selection
INPUT_EXCEL = Path(__file__).parent.parent.parent / "files" / "exams" / "exams-2026-01.xlsm"


# @st.cache_data
def load_data() -> pd.DataFrame:
    """Διαβάζει τα δεδομένα από το Excel (sheet Data)."""
    
    # Έλεγχος ύπαρξης αρχείου
    if not INPUT_EXCEL.exists():
        st.error(f"❌ Το αρχείο {INPUT_EXCEL} δεν βρέθηκε!")
        st.info(f"Αναζητούμενη διαδρομή: {INPUT_EXCEL.absolute()}")
        st.stop()
    
    try:
        # Έλεγχος διαθέσιμων sheets
        excel_file = pd.ExcelFile(INPUT_EXCEL)
        available_sheets = excel_file.sheet_names
        
        if INPUT_SHEET not in available_sheets:
            st.error(f"❌ Το sheet '{INPUT_SHEET}' δεν βρέθηκε στο αρχείο!")
            st.info(f"Διαθέσιμα sheets: {', '.join(available_sheets)}")
            st.stop()
        
        df = pd.read_excel(INPUT_EXCEL, sheet_name=INPUT_SHEET, usecols="A:I")
    except Exception as e:
        st.error(f"❌ Σφάλμα κατά το άνοιγμα του αρχείου: {e}")
        st.stop()

    # Βεβαιώσου ότι τα ονόματα στηλών ταιριάζουν με αυτά
    required_cols = [
        "course_id",
        "course_name",
        "semester",
        "instructor",
        "exam_date",
        "start_time",
        "room",
        "notes",
    ]
    missing = [c for c in required_cols if c not in df.columns]
    if missing:
        raise ValueError(f"Λείπουν οι στήλες: {missing}")

    # Μετατροπές τύπων
    df["exam_date"] = pd.to_datetime(df["exam_date"]).dt.date  # μόνο ημερομηνία
    
    # Drop rows where exam_date is missing
    df = df.dropna(subset=["exam_date"])

    # Add day of week column (in Greek)
    day_names_greek = {
        0: 'Δευτέρα',
        1: 'Τρίτη', 
        2: 'Τετάρτη',
        3: 'Πέμπτη',
        4: 'Παρασκευή',
        5: 'Σάββατο',
        6: 'Κυριακή'
    }
    df["day_of_week"] = pd.to_datetime(df["exam_date"]).dt.dayofweek.map(day_names_greek)

    # Αν start_time είναι string τύπου "09:00"
    df["start_time"] = df["start_time"].astype(str)

    # Συνένωση σε datetime για αρχή
    df["start_dt"] = pd.to_datetime(
        df["exam_date"].astype(str) + " " + df["start_time"],
        errors="coerce",
    )

    # Υπολογισμός end_dt με default διάρκεια
    # Στήλη εβδομάδας (για weekly views)
    df["iso_week_number"] = df["start_dt"].dt.isocalendar().week
    
    # Δημιουργία αντιστοίχησης για εβδομάδες (1, 2, 3, ... αντί για ISO week numbers)
    unique_weeks = sorted(df["iso_week_number"].unique())
    week_mapping = {iso_week: idx + 1 for idx, iso_week in enumerate(unique_weeks)}
    df["week_number"] = df["iso_week_number"].map(week_mapping)

    return df

def reload():
    """Clear cache to force reload from Google Sheets"""
    st.cache_data.clear()


def create_weekly_calendar_document(df: pd.DataFrame, include_epitirites: bool = True) -> bytes:
    """Δημιουργεί Word έγγραφο με εβδομαδιαίο πρόγραμμα εξετάσεων σε μορφή ημερολογίου"""
    doc = Document()
    
    # Ρυθμίσεις σελίδας - landscape για καλύτερη εμφάνιση
    section = doc.sections[0]
    section.orientation = 1  # Landscape
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    
    # Τίτλος
    title = doc.add_heading('Πρόγραμμα Εξετάσεων Χειμερινού Εξαμήνου 2025-2026', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # # Ημερομηνία έκδοσης
    # date_para = doc.add_paragraph(f'Ημερομηνία έκδοσης: {datetime.now().strftime("%d/%m/%Y")}')
    # date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # doc.add_paragraph()
    
    # Ομαδοποίηση κατά εβδομάδα
    weeks = sorted(df['week_number'].unique())
    
    for week_idx, week in enumerate(weeks):
        df_week = df[df['week_number'] == week].sort_values(by=['exam_date', 'start_time'])
        
        if df_week.empty:
            continue
        
        # Εύρεση ημερών της εβδομάδας (εργάσιμες μόνο)
        days = sorted(df_week['exam_date'].unique())
        if not days:
            continue
            
        week_start = days[0]
        week_end = days[-1]
        
        # Επικεφαλίδα εβδομάδας - χρήση ευρετηρίου από 1
        week_heading = doc.add_heading(
            f'Εβδομάδα {week_idx + 1} ({week_start.strftime("%d/%m/%Y")} - {week_end.strftime("%d/%m/%Y")})', 
            level=1
        )
        
        # Δημιουργία πίνακα ημερολογίου: Ώρες x Ημέρες
        time_slots = ['9:00', '12:00', '15:00', '18:00']
        num_days = len(days)
        
        # Πίνακας: 1 σειρά επικεφαλίδας + time slots, 1 στήλη για ώρες + ημέρες
        table = doc.add_table(rows=len(time_slots) + 1, cols=num_days + 1)
        table.style = 'Light Grid Accent 1'
        
        # Επικεφαλίδα - κενό κελί πάνω αριστερά
        table.rows[0].cells[0].text = ''
        
        # Επικεφαλίδες ημερών
        day_names_map = {0: 'Δευ', 1: 'Τρί', 2: 'Τετ', 3: 'Πέμ', 4: 'Παρ', 5: 'Σάβ', 6: 'Κυρ'}
        for day_idx, day in enumerate(days):
            cell = table.rows[0].cells[day_idx + 1]
            day_num = pd.to_datetime(day).dayofweek
            day_name = day_names_map.get(day_num, '')
            cell.text = f'{day_name} {day.strftime("%d/%m")}'
            # Στυλ επικεφαλίδας
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = 'Calibri'
                    run.font.bold = True
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(255, 255, 255)  # Λευκό
        
        # Στήλη ωρών
        for time_idx, time_slot in enumerate(time_slots):
            cell = table.rows[time_idx + 1].cells[0]
            cell.text = time_slot
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = 'Calibri'
                    run.font.bold = True
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(255, 255, 255)  # Λευκό
        
        # Δημιουργία dictionary για συγκέντρωση εξετάσεων ανά κελί
        from collections import defaultdict
        cell_exams = defaultdict(list)
        
        for day_idx, day in enumerate(days):
            df_day = df_week[df_week['exam_date'] == day]
            
            for _, exam in df_day.iterrows():
                # Εύρεση χρονικής θέσης
                exam_time = str(exam['start_time'])
                
                # Εξαγωγή ώρας από το start_time
                if ':' in exam_time:
                    exam_hour = int(exam_time.split(':')[0])
                else:
                    try:
                        exam_hour = int(float(exam_time))
                    except:
                        continue
                
                # Βρες το σωστό time slot
                time_row = None
                for time_idx, time_slot in enumerate(time_slots):
                    slot_hour = int(time_slot.split(':')[0])
                    if exam_hour == slot_hour:
                        time_row = time_idx + 1
                        break
                
                if time_row is None:
                    continue
                
                # Προσθήκη εξέτασης στο dictionary
                cell_key = (time_row, day_idx + 1)
                cell_exams[cell_key].append({
                    'time': exam_time,
                    'semester': f"Εξάμ.{int(exam['semester'])}" if pd.notna(exam['semester']) else '',
                    'course': str(exam['course_name']) if pd.notna(exam['course_name']) else '',
                    'instructor': f'({str(exam['instructor'])})' if pd.notna(exam['instructor']) else '',
                    'room': str(exam['room']) if pd.notna(exam['room']) else '',
                    'epitirites': f'Επιτηρητές: [{str(exam['epitirites'])}]' if pd.notna(exam['epitirites']) else ''
                })
        
        # Συμπλήρωση κελιών με όλες τις εξετάσεις
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        
        # Εφαρμογή χρώματος στη γραμμή επικεφαλίδας (πιο σκούρο)
        for col_idx in range(0, num_days + 1):
            cell = table.rows[0].cells[col_idx]
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), '4472C4')  # Σκούρο μπλε για επικεφαλίδα
            cell._element.get_or_add_tcPr().append(shading_elm)
        
        # Εφαρμογή εναλλασσόμενων ανοιχτών χρωμάτων στις σειρές δεδομένων
        for row_idx in range(1, len(time_slots) + 1):
            # Πολύ ανοιχτά χρώματα για άρτιες και περιττές σειρές
            row_color = 'E7EFF7' if row_idx % 2 == 1 else 'D9E2F3'  # Πολύ ανοιχτό για περιττές, ανοιχτό για άρτιες
            
            for col_idx in range(0, num_days + 1):  # Συμπεριλαμβάνεται και η στήλη ωρών
                cell = table.rows[row_idx].cells[col_idx]
                shading_elm = OxmlElement('w:shd')
                # Χρήση σκούρου μπλε για στήλη ωρών, ανοιχτά χρώματα για υπόλοιπες
                if col_idx == 0:
                    shading_elm.set(qn('w:fill'), '4472C4')  # Σκούρο μπλε για στήλη ωρών
                else:
                    shading_elm.set(qn('w:fill'), row_color)
                cell._element.get_or_add_tcPr().append(shading_elm)
        
        # Προσθήκη περιεχομένου εξετάσεων
        for (time_row, day_col), exams in cell_exams.items():
            cell = table.rows[time_row].cells[day_col]
            
            # Δημιουργία περιεχομένου για όλες τις εξετάσεις
            cell_content = []
            for exam in exams:
                exam_text = f"{exam['semester']} - {exam['course']}\n{exam['instructor']}"
                if exam['room']:
                    exam_text += f"\n{exam['room']}"
                if include_epitirites and exam['epitirites']:
                    exam_text += f"\n{exam['epitirites']}"
                cell_content.append(exam_text)
            
            # Ενωση με διαχωριστικό γραμμή
            cell.text = '\n--------------------------------\n'.join(cell_content)
            
            # Στυλ κελιού
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(8)
        
        # Προσαρμογή πλάτους στηλών
        for row in table.rows:
            row.cells[0].width = Inches(0.6)  # Στήλη ωρών
            for i in range(1, num_days + 1):
                row.cells[i].width = Inches(2.0)  # Στήλες ημερών
        
        if week_idx < len(weeks) - 1:
            doc.add_page_break()
    
    # Αποθήκευση σε buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


tab_full_table, tab_instructor_filter, tab_semester_filter, tab_epitiritis_filter, tab_calendar, tab_export_weekly = st.tabs(
    [
        "Πλήρης Πίνακας Εξετάσεων",
        "Φιλτράρισμα κατά Διδάσκοντα",
        "Φιλτράρισμα κατά Εξάμηνο",
        "Φιλτράρισμα κατά Επιτηρητή",
        "Ημερολόγιο Εξετάσεων",
        "Εξαγωγή Εβδομαδιαίου Προγράμματος"
    ]
)    

df = load_data()


with tab_full_table:
    st.subheader("Πλήρης Πίνακας Εξετάσεων")
    # Απόκρυψη βοηθητικών στηλών
    display_cols = [col for col in df.columns if col not in ['start_dt', 'iso_week_number']]
    st.dataframe(df[display_cols])

instructors = sorted(df["instructor"].unique().tolist())


with tab_instructor_filter:
    selected_instructor = st.selectbox(
        "Επιλέξτε διδάσκοντα για φιλτράρισμα:",
        options=instructors)

    df_instr = df[df["instructor"] == selected_instructor].sort_values(
        by=["start_dt"]
    )  

    st.subheader(f"Πρόγραμμα Εξετάσεων Διδάσκοντα - {selected_instructor}")
    # Απόκρυψη βοηθητικών στηλών
    display_cols = [col for col in df_instr.columns if col not in ['start_dt', 'iso_week_number']]
    st.dataframe(df_instr[display_cols])

with tab_semester_filter:
    semesters = sorted(df["semester"].unique().tolist())
    selected_semester = st.selectbox(
        "Επιλέξτε εξάμηνο για φιλτράρισμα:",
        options=semesters
    )

    df_sem = df[df["semester"] == selected_semester].sort_values(
        by=["start_dt"]
    )  

    st.subheader(f"Πρόγραμμα Εξετάσεων Εξαμήνου - {selected_semester}")
    # Απόκρυψη βοηθητικών στηλών
    display_cols = [col for col in df_sem.columns if col not in ['start_dt', 'iso_week_number']]
    st.dataframe(df_sem[display_cols], height=600)    

with tab_epitiritis_filter:
    # Get unique epitirites values, handling NaN and splitting comma-separated values
    epitirites_list = []
    for val in df["epitirites"].dropna().unique():
        # Split by comma if multiple epitirites per exam
        if pd.notna(val):
            epitirites_list.extend([e.strip() for e in str(val).split(',')])
    
    # Get unique and sorted list
    epitirites_unique = sorted(list(set(epitirites_list)))
    
    if epitirites_unique:
        selected_epitiritis = st.selectbox(
            "Επιλέξτε επιτηρητή για φιλτράρισμα:",
            options=epitirites_unique
        )
        
        # Filter rows where selected epitiritis appears (handling comma-separated values)
        df_epit = df[df["epitirites"].apply(
            lambda x: selected_epitiritis in str(x) if pd.notna(x) else False
        )].sort_values(by=["start_dt"])
        
        st.subheader(f"Πρόγραμμα Επιτηρήσεων - {selected_epitiritis}")
        
        # Απόκρυψη βοηθητικών στηλών
        display_cols = [col for col in df_epit.columns if col not in ['start_dt', 'iso_week_number']]
        st.dataframe(df_epit[display_cols], height=600)
        
        # Statistics
        st.markdown("### Στατιστικά")
        col1, col2 = st.columns(2)
        with col1:
            st.metric("Συνολικές Επιτηρήσεις", len(df_epit))
        with col2:
            unique_weeks = df_epit['week_number'].nunique()
            st.metric("Εβδομάδες με Επιτηρήσεις", unique_weeks)
    else:
        st.warning("⚠️ Δεν βρέθηκαν δεδομένα επιτηρητών στο αρχείο.")

with tab_calendar:
    st.subheader("Ημερολόγιο Εξετάσεων")

    # Add semester filter for calendar
    semesters_all = sorted(df["semester"].unique().tolist())
    semester_options = [f"Εξάμηνο {int(s)}" for s in semesters_all]

    selected_calendar_semesters = st.multiselect(
        "Φιλτράρισμα ημερολογίου κατά εξάμηνο:",
        options=semester_options,
        default=semester_options,  # All semesters selected by default
        key="calendar_semester_filter"
    )

    # Filter dataframe based on selection
    if not selected_calendar_semesters or len(selected_calendar_semesters) == len(semester_options):
        df_calendar = df
    else:
        # Extract semester numbers from "Εξάμηνο X"
        semester_nums = [int(s.split()[-1]) for s in selected_calendar_semesters]
        df_calendar = df[df["semester"].isin(semester_nums)]

    # Get the earliest exam date for initial calendar view
    initial_date = df_calendar["exam_date"].min() if not df_calendar.empty else datetime.now().date()

    calendar_options = {
        "initialView": "timeGridWeek",
        "initialDate": initial_date.strftime("%Y-%m-%d"),
        "selectable": True,
        "weekends": False,
        "slotMinTime": "08:00:00",
        "slotMaxTime": "22:00:00",
        "headerToolbar": {
            "left": "today prev,next",
            "center": "title",
            "right": "dayGridMonth,timeGridWeek,timeGridDay"
        }
    }

    # Χρώματα ανά εξάμηνο
    semester_colors = {
        1: '#E74C3C',  # Κόκκινο
        2: '#3498DB',  # Μπλε
        3: '#2ECC71',  # Πράσινο
        4: '#F39C12',  # Πορτοκαλί
        5: '#9B59B6',  # Μωβ
        6: '#1ABC9C',  # Τιρκουάζ
        7: '#E67E22',  # Πορτοκαλί σκούρο
        8: '#34495E',  # Γκρι-μπλε
        9: '#16A085',  # Πράσινο-μπλε
        10: '#D35400', # Καφέ-πορτοκαλί
    }
    
    # Convert exam data to calendar events
    calendar_events = []
    for _, row in df_calendar.iterrows():
        if pd.notna(row["start_dt"]):
            # Format as string YYYY-MM-DDTHH:MM:SS
            start_str = row["start_dt"].strftime("%Y-%m-%dT%H:%M:%S")
            
            # Calculate end time (2 hours after start)
            end_dt = row["start_dt"] + timedelta(hours=2)
            end_str = end_dt.strftime("%Y-%m-%dT%H:%M:%S")
            
            # Safely handle potential None values, convert to string, and remove problematic characters
            def clean_text(value):
                if pd.notna(value):
                    # Convert to string and remove newlines, quotes, backslashes
                    text = str(value).replace('\n', ' ').replace('\r', ' ')
                    text = text.replace('"', '').replace("'", '').replace('\\', '')
                    return text.strip()
                return ""
            
            course_name = clean_text(row['course_name'])
            instructor = clean_text(row['instructor'])
            room = clean_text(row['room'])
            semester = int(row['semester']) if pd.notna(row['semester']) else 1
            semester_str = str(semester)
            
            # Επιλογή χρώματος βάσει εξαμήνου
            color = semester_colors.get(semester, '#95A5A6')  # Default γκρι αν δεν βρεθεί
            
            event = {
                "title": f'Εξ.{semester_str} - {course_name} - {instructor}',
                "start": start_str,
                "end": end_str,
                "color": color
            }
            calendar_events.append(event)

    # Debug: show number of events
    if not selected_calendar_semesters or len(selected_calendar_semesters) == len(semester_options):
        st.write(f"📅 Σύνολο εξετάσεων: {len(calendar_events)} (όλα τα εξάμηνα)")
    else:
        semesters_text = ", ".join(selected_calendar_semesters)
        st.write(f"📅 Σύνολο εξετάσεων: {len(calendar_events)} ({semesters_text})")

    # Initialize session state for calendar display
    if 'show_exam_calendar' not in st.session_state:
        st.session_state.show_exam_calendar = False

    # Button to show calendar
    if not st.session_state.show_exam_calendar:
        if st.button("📅 Εμφάνιση Ημερολογίου", key="show_exam_cal_btn"):
            st.session_state.show_exam_calendar = True
            st.rerun()
    
    # Render calendar if button was clicked
    if st.session_state.show_exam_calendar:
        if calendar_events:
            calendar_data = calendar(
                events=calendar_events,
                options=calendar_options
            )
        else:
            st.info("Δεν υπάρχουν εξετάσεις για εμφάνιση με τα επιλεγμένα φίλτρα.")

# st.write("Calendar interaction information:", calendar_data)


with tab_export_weekly:
    st.subheader("Εξαγωγή Εβδομαδιαίου Προγράμματος Εξετάσεων")
    st.markdown("Δημιουργήστε αρχείο Word με το εβδομαδιαίο πρόγραμμα εξετάσεων για διανομή σε συναδέλφους.")
    
    # Φίλτρα για εξαγωγή
    st.markdown("### Επιλογές Φιλτραρίσματος")
    
    col1, col2 = st.columns(2)
    
    with col1:
        semesters_export = sorted(df["semester"].unique().tolist())
        semester_options_export = [f"Εξάμηνο {int(s)}" for s in semesters_export]
        
        selected_export_semesters = st.multiselect(
            "Επιλέξτε εξάμηνα:",
            options=semester_options_export,
            default=semester_options_export,
            key="export_semester_filter"
        )
    
    with col2:
        # Φίλτρο εβδομάδων
        weeks_available = sorted(df['week_number'].unique().tolist())
        week_options = [f"Εβδομάδα {int(w)}" for w in weeks_available]
        
        selected_export_weeks = st.multiselect(
            "Επιλέξτε εβδομάδες:",
            options=week_options,
            default=week_options,
            key="export_week_filter"
        )
    
    # Checkbox για επιτηρητές
    include_epitirites = st.checkbox(
        "Συμπερίληψη επιτηρητών στο αρχείο Word",
        value=True,
        key="include_epitirites_checkbox",
        help="Επιλέξτε αν θέλετε να περιλαμβάνονται οι επιτηρητές στο εξαγόμενο αρχείο"
    )
    
    # Φιλτράρισμα δεδομένων
    df_export = df.copy()
    
    if selected_export_semesters and len(selected_export_semesters) < len(semester_options_export):
        semester_nums = [int(s.split()[-1]) for s in selected_export_semesters]
        df_export = df_export[df_export["semester"].isin(semester_nums)]
    
    if selected_export_weeks and len(selected_export_weeks) < len(week_options):
        week_nums = [int(w.split()[-1]) for w in selected_export_weeks]
        df_export = df_export[df_export["week_number"].isin(week_nums)]
    
    # Προεπισκόπηση
    st.markdown("### Προεπισκόπηση Δεδομένων")
    st.write(f"Σύνολο εξετάσεων προς εξαγωγή: {len(df_export)}")
    
    if not df_export.empty:
        st.dataframe(
            df_export[['exam_date', 'day_of_week', 'start_time', 'semester', 
                       'course_name', 'instructor', 'room', 'epitirites']].sort_values(by=['exam_date', 'start_time']),
            height=400
        )
        
        # Κουμπί λήψης
        st.markdown("### Λήψη Αρχείου")
        
        try:
            word_file = create_weekly_calendar_document(df_export, include_epitirites=include_epitirites)
            
            # Όνομα αρχείου με ημερομηνία
            filename = f"Πρόγραμμα_Εξετάσεων_{program_selection}_{exam_period}.docx"
            
            st.download_button(
                label="📥 Λήψη Word Αρχείου",
                data=word_file,
                file_name=filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                help="Κατεβάστε το εβδομαδιαίο πρόγραμμα εξετάσεων σε μορφή Word"
            )
            
            st.success("✅ Το αρχείο είναι έτοιμο για λήψη!")
            
        except Exception as e:
            st.error(f"Σφάλμα κατά τη δημιουργία του αρχείου: {e}")
    else:
        st.warning("⚠️ Δεν υπάρχουν δεδομένα με τα επιλεγμένα φίλτρα.")