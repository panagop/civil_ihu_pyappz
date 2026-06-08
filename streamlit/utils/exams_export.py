import io
from collections import defaultdict

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def create_weekly_calendar_document(
    df: pd.DataFrame,
    period: str,
    include_epitirites: bool = True,
) -> bytes:
    """Δημιουργεί Word έγγραφο με εβδομαδιαίο πρόγραμμα εξετάσεων σε μορφή ημερολογίου."""
    doc = Document()

    section = doc.sections[0]
    section.orientation = 1  # Landscape
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)

    title = doc.add_heading(f'Πρόγραμμα Εξετάσεων {period} Εξάμηνο 2025-2026', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    weeks = sorted(df['week_number'].unique())

    for week_idx, week in enumerate(weeks):
        df_week = df[df['week_number'] == week].sort_values(by=['exam_date', 'start_time'])

        if df_week.empty:
            continue

        days = sorted(df_week['exam_date'].unique())
        if not days:
            continue

        week_start = days[0]
        week_end = days[-1]

        doc.add_heading(
            f'Εβδομάδα {week_idx + 1} ({week_start.strftime("%d/%m/%Y")} - {week_end.strftime("%d/%m/%Y")})',
            level=1,
        )

        time_slots = ['9:00', '12:00', '15:00', '18:00']
        num_days = len(days)

        table = doc.add_table(rows=len(time_slots) + 1, cols=num_days + 1)
        table.style = 'Light Grid Accent 1'

        table.rows[0].cells[0].text = ''

        day_names_map = {0: 'Δευ', 1: 'Τρί', 2: 'Τετ', 3: 'Πέμ', 4: 'Παρ', 5: 'Σάβ', 6: 'Κυρ'}
        for day_idx, day in enumerate(days):
            cell = table.rows[0].cells[day_idx + 1]
            day_num = pd.to_datetime(day).dayofweek
            day_name = day_names_map.get(day_num, '')
            cell.text = f'{day_name} {day.strftime("%d/%m")}'
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = 'Calibri'
                    run.font.bold = True
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(255, 255, 255)

        for time_idx, time_slot in enumerate(time_slots):
            cell = table.rows[time_idx + 1].cells[0]
            cell.text = time_slot
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = 'Calibri'
                    run.font.bold = True
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(255, 255, 255)

        cell_exams = defaultdict(list)

        for day_idx, day in enumerate(days):
            df_day = df_week[df_week['exam_date'] == day]

            for _, exam in df_day.iterrows():
                exam_time = str(exam['start_time'])

                if ':' in exam_time:
                    exam_hour = int(exam_time.split(':')[0])
                else:
                    try:
                        exam_hour = int(float(exam_time))
                    except Exception:
                        continue

                time_row = None
                for time_idx, time_slot in enumerate(time_slots):
                    slot_hour = int(time_slot.split(':')[0])
                    if exam_hour == slot_hour:
                        time_row = time_idx + 1
                        break

                if time_row is None:
                    continue

                cell_key = (time_row, day_idx + 1)
                cell_exams[cell_key].append({
                    'time': exam_time,
                    'semester': f"Εξάμ.{int(exam['semester'])}" if pd.notna(exam['semester']) else '',
                    'course': str(exam['course_name']) if pd.notna(exam['course_name']) else '',
                    'instructor': f'({str(exam["instructor"])})' if pd.notna(exam['instructor']) else '',
                    'room': str(exam['room']) if pd.notna(exam['room']) else '',
                    'epitirites': f'Επιτηρητές: [{str(exam["epitirites"])}]' if pd.notna(exam['epitirites']) else '',
                })

        for col_idx in range(0, num_days + 1):
            cell = table.rows[0].cells[col_idx]
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), '4472C4')
            cell._element.get_or_add_tcPr().append(shading_elm)

        for row_idx in range(1, len(time_slots) + 1):
            row_color = 'E7EFF7' if row_idx % 2 == 1 else 'D9E2F3'

            for col_idx in range(0, num_days + 1):
                cell = table.rows[row_idx].cells[col_idx]
                shading_elm = OxmlElement('w:shd')
                if col_idx == 0:
                    shading_elm.set(qn('w:fill'), '4472C4')
                else:
                    shading_elm.set(qn('w:fill'), row_color)
                cell._element.get_or_add_tcPr().append(shading_elm)

        for (time_row, day_col), exams in cell_exams.items():
            cell = table.rows[time_row].cells[day_col]

            cell_content = []
            for exam in exams:
                exam_text = f"{exam['semester']} - {exam['course']}\n{exam['instructor']}"
                if exam['room']:
                    exam_text += f"\n{exam['room']}"
                if include_epitirites and exam['epitirites']:
                    exam_text += f"\n{exam['epitirites']}"
                cell_content.append(exam_text)

            cell.text = '\n--------------------------------\n'.join(cell_content)

            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(8)

        for row in table.rows:
            row.cells[0].width = Inches(0.6)
            for i in range(1, num_days + 1):
                row.cells[i].width = Inches(2.0)

        if week_idx < len(weeks) - 1:
            doc.add_page_break()

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
