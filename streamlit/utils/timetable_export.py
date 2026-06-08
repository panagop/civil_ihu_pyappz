import io
from collections import defaultdict

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def create_weekly_timetable_document(df: pd.DataFrame, period: str) -> bytes:
    """Δημιουργεί Word έγγραφο με εβδομαδιαίο πρόγραμμα μαθημάτων."""
    doc = Document()

    section = doc.sections[0]
    section.orientation = 1  # Landscape
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)

    title = doc.add_heading(
        f'Εβδομαδιαίο Πρόγραμμα Μαθημάτων - {period} Εξάμηνο 2025-2026', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    semesters = sorted(df['semester'].unique())

    for sem_idx, semester in enumerate(semesters):
        df_sem = df[df['semester'] == semester]

        if df_sem.empty:
            continue

        doc.add_heading(f'Εξάμηνο {int(semester)}', level=1)

        time_slots = list(range(9, 22))
        time_slots_str = [f"{h}:00" for h in time_slots]

        day_names = ['Δευτέρα', 'Τρίτη', 'Τετάρτη', 'Πέμπτη', 'Παρασκευή']

        cell_classes = defaultdict(list)

        for day_idx, day_name in enumerate(day_names):
            df_day = df_sem[df_sem['day'] == day_name]

            for _, class_row in df_day.iterrows():
                start_time_str = str(class_row['start_time'])
                if ':' in start_time_str:
                    class_hour = int(start_time_str.split(':')[0])
                else:
                    try:
                        class_hour = int(float(start_time_str))
                    except Exception:
                        continue

                try:
                    time_idx = time_slots.index(class_hour)
                except ValueError:
                    continue

                cell_key = (time_idx, day_idx)
                cell_classes[cell_key].append({
                    'course': str(class_row['full_class_name']) if pd.notna(class_row['full_class_name']) else '',
                    'instructor': str(class_row['instructors']) if pd.notna(class_row['instructors']) else '',
                    'room': str(class_row['room']) if pd.notna(class_row['room']) else '',
                    'duration': int(class_row['duration']) if pd.notna(class_row['duration']) else 1,
                })

        max_simultaneous = 1
        for classes in cell_classes.values():
            max_simultaneous = max(max_simultaneous, len(classes))

        total_cols = 1 + (len(day_names) * max_simultaneous)
        table = doc.add_table(rows=len(time_slots) + 1, cols=total_cols)
        table.style = 'Light Grid Accent 1'

        table.rows[0].cells[0].text = 'Ώρα'

        for day_idx, day_name in enumerate(day_names):
            start_col = 1 + (day_idx * max_simultaneous)
            cell = table.rows[0].cells[start_col]

            if max_simultaneous > 1:
                for sub_col in range(1, max_simultaneous):
                    cell.merge(table.rows[0].cells[start_col + sub_col])

            cell.text = day_name
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = 'Calibri'
                    run.font.bold = True
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(255, 255, 255)

        processed_cells = {}

        for time_idx, time_slot in enumerate(time_slots_str):
            row_idx = time_idx + 1

            cell = table.rows[row_idx].cells[0]
            cell.text = time_slot
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = 'Calibri'
                    run.font.bold = True
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(255, 255, 255)

            for day_idx in range(len(day_names)):
                cell_key = (time_idx, day_idx)
                classes = cell_classes.get(cell_key, [])

                start_col = 1 + (day_idx * max_simultaneous)

                if len(classes) <= 1:
                    if (time_idx, day_idx, 0) in processed_cells:
                        continue

                    if len(classes) == 1:
                        cls = classes[0]
                        duration = cls['duration']

                        for dur in range(duration):
                            for sub_col_idx in range(max_simultaneous):
                                processed_cells[(time_idx + dur, day_idx, sub_col_idx)] = True

                        cells_to_merge = []
                        for dur_offset in range(duration):
                            if time_idx + dur_offset < len(time_slots):
                                for sub_col_offset in range(max_simultaneous):
                                    cells_to_merge.append(
                                        table.rows[row_idx + dur_offset].cells[start_col + sub_col_offset])

                        cell = cells_to_merge[0]

                        for merge_cell in cells_to_merge[1:]:
                            if merge_cell != cell:
                                try:
                                    cell.merge(merge_cell)
                                except Exception:
                                    pass

                        class_text = f"{cls['course']}\n{cls['instructor']}"
                        if cls['room']:
                            class_text += f"\n{cls['room']}"
                        cell.text = class_text

                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for run in paragraph.runs:
                                run.font.name = 'Calibri'
                                run.font.size = Pt(8)
                    else:
                        for sub_col_idx in range(max_simultaneous):
                            processed_cells[(time_idx, day_idx, sub_col_idx)] = True

                        if max_simultaneous > 1:
                            cell = table.rows[row_idx].cells[start_col]
                            for sub_col_offset in range(1, max_simultaneous):
                                next_col_cell = table.rows[row_idx].cells[start_col + sub_col_offset]
                                try:
                                    cell.merge(next_col_cell)
                                except Exception:
                                    pass
                else:
                    for cls_idx, cls in enumerate(classes[:max_simultaneous]):
                        if (time_idx, day_idx, cls_idx) in processed_cells:
                            continue

                        col = start_col + cls_idx
                        duration = cls['duration']

                        for dur in range(duration):
                            processed_cells[(time_idx + dur, day_idx, cls_idx)] = True

                        cell = table.rows[row_idx].cells[col]

                        if duration > 1:
                            for dur_offset in range(1, duration):
                                if time_idx + dur_offset < len(time_slots):
                                    next_cell = table.rows[row_idx + dur_offset].cells[col]
                                    try:
                                        cell.merge(next_cell)
                                    except Exception:
                                        pass

                        class_text = f"{cls['course']}\n{cls['instructor']}"
                        if cls['room']:
                            class_text += f"\n{cls['room']}"
                        cell.text = class_text

                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for run in paragraph.runs:
                                run.font.name = 'Calibri'
                                run.font.size = Pt(8)

        for col_idx in range(0, total_cols):
            cell = table.rows[0].cells[col_idx]
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), '4472C4')
            cell._element.get_or_add_tcPr().append(shading_elm)

        for row_idx in range(1, len(time_slots) + 1):
            for col_idx in range(0, total_cols):
                cell = table.rows[row_idx].cells[col_idx]
                shading_elm = OxmlElement('w:shd')

                if col_idx == 0:
                    shading_elm.set(qn('w:fill'), '4472C4')
                else:
                    cell_text = cell.text.strip()
                    if cell_text:
                        shading_elm.set(qn('w:fill'), 'E7EFF7')
                    else:
                        shading_elm.set(qn('w:fill'), 'FFFFFF')

                cell._element.get_or_add_tcPr().append(shading_elm)

        for row in table.rows:
            row.height = Inches(0.4)
            row.cells[0].width = Inches(0.7)
            for i in range(1, len(day_names) * max_simultaneous + 1):
                row.cells[i].width = Inches(1.5)

        if sem_idx < len(semesters) - 1:
            doc.add_page_break()

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
